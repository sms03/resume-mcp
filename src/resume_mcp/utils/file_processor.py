"""
File processing utilities for handling PDF and DOCX resume uploads.
"""

import io
import tempfile
from pathlib import Path
from typing import Union, Optional, BinaryIO
import logging

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False


logger = logging.getLogger(__name__)


class FileProcessor:
    """Handles extraction of text from various file formats."""
    
    def __init__(self):
        """Initialize file processor with available libraries."""
        self.supported_formats = []
        
        if PDF_AVAILABLE or PYMUPDF_AVAILABLE:
            self.supported_formats.extend(['.pdf'])
        
        if DOCX_AVAILABLE:
            self.supported_formats.extend(['.docx', '.doc'])
        
        # Always support plain text
        self.supported_formats.extend(['.txt', '.text'])
        
        logger.info(f"FileProcessor initialized with support for: {self.supported_formats}")
    
    def is_supported(self, filename: str) -> bool:
        """Check if file format is supported."""
        file_ext = Path(filename).suffix.lower()
        return file_ext in self.supported_formats
    
    def extract_text(self, file_content: Union[bytes, BinaryIO], filename: str) -> str:
        """
        Extract text from uploaded file.
        
        Args:
            file_content: File content as bytes or file-like object
            filename: Original filename to determine format
            
        Returns:
            Extracted text content
            
        Raises:
            ValueError: If file format is not supported
            Exception: If text extraction fails
        """
        file_ext = Path(filename).suffix.lower()
        
        if not self.is_supported(filename):
            raise ValueError(f"Unsupported file format: {file_ext}")
        
        # Convert to bytes if needed
        if hasattr(file_content, 'read'):
            content_bytes = file_content.read()
        else:
            content_bytes = file_content
        
        try:
            if file_ext == '.pdf':
                return self._extract_pdf_text(content_bytes)
            elif file_ext in ['.docx', '.doc']:
                return self._extract_docx_text(content_bytes)
            elif file_ext in ['.txt', '.text']:
                return self._extract_text_file(content_bytes)
            else:
                raise ValueError(f"No handler for file type: {file_ext}")
        
        except Exception as e:
            logger.error(f"Failed to extract text from {filename}: {e}")
            raise Exception(f"Text extraction failed: {str(e)}")
    
    def _extract_pdf_text(self, content_bytes: bytes) -> str:
        """Extract text from PDF file."""
        text_parts = []
        
        # Try PyMuPDF first (generally better)
        if PYMUPDF_AVAILABLE:
            try:
                doc = fitz.open(stream=content_bytes, filetype="pdf")
                for page_num in range(doc.page_count):
                    page = doc[page_num]
                    text_parts.append(page.get_text())
                doc.close()
                return "\n".join(text_parts)
            except Exception as e:
                logger.warning(f"PyMuPDF extraction failed: {e}, trying PyPDF2")
        
        # Fallback to PyPDF2
        if PDF_AVAILABLE:
            try:
                pdf_file = io.BytesIO(content_bytes)
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                
                for page in pdf_reader.pages:
                    text_parts.append(page.extract_text())
                
                return "\n".join(text_parts)
            except Exception as e:
                logger.error(f"PyPDF2 extraction failed: {e}")
                raise
        
        raise Exception("No PDF processing library available")
    
    def _extract_docx_text(self, content_bytes: bytes) -> str:
        """Extract text from DOCX file."""
        if not DOCX_AVAILABLE:
            raise Exception("python-docx library not available")
        
        try:
            # Create a temporary file to work with python-docx
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                temp_file.write(content_bytes)
                temp_path = temp_file.name
            
            try:
                doc = Document(temp_path)
                text_parts = []
                
                # Extract text from paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text)
                
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                text_parts.append(cell.text)
                
                return "\n".join(text_parts)
            
            finally:
                # Clean up temporary file
                try:
                    Path(temp_path).unlink()
                except:
                    pass
        
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")
    
    def _extract_text_file(self, content_bytes: bytes) -> str:
        """Extract text from plain text file."""
        try:
            # Try UTF-8 first
            return content_bytes.decode('utf-8')
        except UnicodeDecodeError:
            try:
                # Fallback to latin-1
                return content_bytes.decode('latin-1')
            except UnicodeDecodeError:
                # Last resort: ignore errors
                return content_bytes.decode('utf-8', errors='ignore')
    
    def validate_file_size(self, content_bytes: bytes, max_size_mb: int = 10) -> bool:
        """
        Validate file size is within limits.
        
        Args:
            content_bytes: File content
            max_size_mb: Maximum size in megabytes
            
        Returns:
            True if file size is acceptable
        """
        size_mb = len(content_bytes) / (1024 * 1024)
        return size_mb <= max_size_mb
    
    def get_file_info(self, content_bytes: bytes, filename: str) -> dict:
        """
        Get file information including size and format.
        
        Args:
            content_bytes: File content
            filename: Original filename
            
        Returns:
            Dictionary with file information
        """
        return {
            'filename': filename,
            'size_bytes': len(content_bytes),
            'size_mb': round(len(content_bytes) / (1024 * 1024), 2),
            'format': Path(filename).suffix.lower(),
            'is_supported': self.is_supported(filename)
        }
    
    @classmethod
    def get_installation_instructions(cls) -> dict:
        """Get instructions for installing missing dependencies."""
        instructions = {
            'pdf_support': {
                'available': PDF_AVAILABLE or PYMUPDF_AVAILABLE,
                'recommended': 'pip install PyMuPDF',
                'alternative': 'pip install PyPDF2'
            },
            'docx_support': {
                'available': DOCX_AVAILABLE,
                'install': 'pip install python-docx'
            }
        }
        
        return instructions


# Global file processor instance
file_processor = FileProcessor()
